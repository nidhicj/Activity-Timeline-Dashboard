from docx import Document
from docx.shared import Inches

# Create the document
doc = Document()
doc.add_heading('Solution Design Document (SDD)', 0)

# Document Details
doc.add_heading('1. Document Details', level=1)
doc.add_table(rows=6, cols=2).style = 'Table Grid'
table = doc.tables[0]
details = [
    ("Project Name", "Activity Timeline Dashboard"),
    ("Author", "Nidhi"),
    ("Date", "2025-06-17"),
    ("Version", "1.0"),
    ("Reviewers", "TBD"),
    ("Status", "Draft")
]
for i, (k, v) in enumerate(details):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# Purpose
doc.add_heading('2. Purpose', level=1)
doc.add_paragraph("""To design and implement a web application that:
- Fetches data from two external APIs at the end of every day
- Aggregates this data based on a common identifier
- Stores the aggregated data in a database
- Displays this data on the website
- Maintains logs for each operation""")

# Scope
doc.add_heading('3. Scope', level=1)
doc.add_paragraph("Included:")
doc.add_paragraph("""
- Backend service to call APIs
- Data aggregation logic
- Database storage
- Logging
- Web UI to display the data""", style='List Bullet')
doc.add_paragraph("Not Included:")
doc.add_paragraph("""
- User authentication (unless specified later)
- Admin panel for data manipulation""", style='List Bullet')

# Assumptions
doc.add_heading('4. Assumptions', level=1)
doc.add_paragraph("""
- APIs are reliable and accessible
- Data returned is well-structured JSON
- API rate limits and authentication (if any) are manageable
- Website is publicly or internally hosted""")

# Architecture
doc.add_heading('5. High-Level Architecture', level=1)
doc.add_paragraph("""
Scheduler → API Fetcher → External APIs → Data Aggregator → Logger + Database → Web Frontend""")

# Data Flow
doc.add_heading('6. Data Flow', level=1)
doc.add_paragraph("""
1. Daily Trigger (e.g., at 11:59 PM): Calls fetch_data()
2. Fetch activityTimeline API and JIRA API JSON responses
3. Parse and merge on id → create a combined DataFrame
4. Save DataFrame as a table to the DB
5. Log: Timestamp, API status, merge status, errors
6. Frontend reads DB and renders data""")

# Tech Stack
doc.add_heading('7. Technology Stack', level=1)
doc.add_paragraph("""
Frontend: React.js
Backend: FastAPI (Python)
Scheduler: APScheduler or Celery
Database: PostgreSQL (or SQLite)
ORM: SQLAlchemy
Hosting: Render / Vercel / VPS
Logs: logging (Python) + DB
VCS: Git + GitHub""")

# DB Schema
doc.add_heading('8. Database Schema (Sample)', level=1)
doc.add_paragraph("""
TABLE aggregated_data (
    id TEXT PRIMARY KEY,
    attr1_from_api1 TEXT,
    attr2_from_api1 TEXT,
    attrX_from_api2 TEXT,
    fetch_timestamp TIMESTAMP
)""")

# Logging
doc.add_heading('9. Logging Schema', level=1)
doc.add_paragraph("""
TABLE logs (
    id SERIAL PRIMARY KEY,
    timestamp TIMESTAMP,
    status TEXT,
    message TEXT,
    level TEXT
)""")

# Security
doc.add_heading('10. Security & Error Handling', level=1)
doc.add_paragraph("""
- API errors (timeouts, bad responses) logged with timestamp
- Retry logic for failed API calls (up to 3 times)
- Sanitization before DB insert
- Display fallback message if DB is empty""")

# Jira Plan
doc.add_heading('11. JIRA Implementation Plan', level=1)
doc.add_paragraph("""Epic: Activity Timeline Dashboard""")
jira_stories = [
    "Set up FastAPI backend and endpoints",
    "Implement scheduler to trigger API calls",
    "Write logic to fetch data from activityTimeline API",
    "Write logic to fetch data from JIRA API",
    "Merge data on `id` and convert to DataFrame",
    "Store data into PostgreSQL DB",
    "Implement logging system",
    "Build React frontend to display data",
    "Connect frontend to backend (REST API)",
    "Write test cases",
    "Deploy application",
    "Document everything in Confluence/Notion"
]
for story in jira_stories:
    doc.add_paragraph(story, style='List Bullet')

# Risks
doc.add_heading('12. Risks', level=1)
doc.add_paragraph("""
- API outages
- Inconsistent schema from external APIs
- DB performance if dataset grows large
- Timezone handling for "end of day"
""")

# Enhancements
doc.add_heading('13. Future Enhancements', level=1)
doc.add_paragraph("""
- Add data visualization (charts, trends)
- Add authentication
- Export data (CSV/Excel)
- API health monitoring dashboard
""")

# Save the file
file_path = "ActivityTimeline_SDD.docx"
doc.save(file_path)
file_path
